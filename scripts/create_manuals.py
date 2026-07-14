"""Build the JC-App operating manual set as editable DOCX files.

The source intentionally contains no credential values, member emails, or
tokens. PDF conversion and visual QA are performed separately.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "manuals"
TODAY = "2026年7月14日"
ORG = "一般社団法人 玉島青年会議所"
BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
GOLD = "FFF4CC"
RED = "FDE9E7"
GRAY = "D9E2F3"
FONT = "Yu Gothic"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, color="667085")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document(doc, manual_name):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11, "1F3A5F", 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(f"玉島青年会議所 会員管理アプリ | {manual_name}")
    set_run_font(r, 8.5, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Ver.1.0 | ")
    set_run_font(r, 8.5, color="667085")
    add_page_field(footer)

    # Ask Word to update the TOC field when the document is opened.
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings.element.append(update)


def add_text(doc, text, bold=False, color=None, align=None, size=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, 10.0)


def add_steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, 10.0)


def add_note(doc, label, text, kind="notice"):
    fill = {"notice": LIGHT_BLUE, "important": GOLD, "warning": RED, "confirm": LIGHT_GRAY}[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, 9.5, bold=True, color=BLUE if kind != "warning" else "9B1C1C")
    r = p.add_run(text)
    set_run_font(r, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_placeholder(doc, description):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E5E7EB")
    set_cell_margins(cell, top=280, bottom=280)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ここに画面キャプチャを挿入\n")
    set_run_font(r, 10, bold=True, color="6B7280")
    r = p.add_run(description)
    set_run_font(r, 8.5, color="6B7280")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_fixed(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_col_width(cells[0], Inches(1.45))
        set_col_width(cells[1], Inches(5.15))
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, 9.2, bold=True, color=BLUE)
        p = cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(table)
    if widths is None:
        widths = [6.6 / len(headers)] * len(headers)
    for index, heading in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_col_width(cell, Inches(widths[index]))
        set_cell_margins(cell)
        set_cell_shading(cell, GRAY)
        p = cell.paragraphs[0]
        r = p.add_run(heading)
        set_run_font(r, 8.7, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_col_width(cells[index], Inches(widths[index]))
            set_cell_margins(cells[index])
            p = cells[index].paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, 8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc, manual_name, subtitle):
    doc.add_paragraph().paragraph_format.space_after = Pt(55)
    add_text(doc, "玉島青年会議所 会員管理アプリ", bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, after=18)
    add_text(doc, manual_name, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER, size=27, after=12)
    add_text(doc, subtitle, color="475467", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=46)
    add_kv_table(doc, [
        ("バージョン", "Ver.1.0"),
        ("作成日", TODAY),
        ("管理部署", ORG),
        ("取扱区分", "内部運用資料。個人情報を含む画面キャプチャは必要最小限にしてください。"),
    ])
    add_note(doc, "重要", "初期パスワード、Secret key、認証トークン、実際の会員メールアドレスは、この文書へ記載しません。", "warning")
    doc.add_page_break()


def add_revision_and_toc(doc, entries):
    doc.add_heading("更新履歴", level=1)
    add_matrix(doc, ["版", "日付", "更新内容", "更新者"], [
        ["1.0", TODAY, "Ver.1 運用開始用に作成", "一般社団法人 玉島青年会議所"],
        ["", "", "", ""],
        ["", "", "", ""],
    ], [0.65, 1.15, 3.8, 1.0])
    add_note(doc, "運用", "担当交代時は、変更箇所をこの表へ追記し、Wordの目次フィールドを更新してから配布用PDFを作り直してください。", "notice")
    doc.add_page_break()
    doc.add_heading("目次", level=1)
    add_text(doc, "Wordでこの文書を開いた際に目次を更新すると、見出しとページ番号が再計算されます。PDFは配布前に更新後のWordから出力してください。", color="475467", size=9.2)
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(entry)
        set_run_font(r, 10.0, color="1F3A5F")
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_page_break()


def add_chapter(doc, number, title, owner, timing, purpose, blocks, placeholder, checkpoints, faq=None, force_break=True):
    if force_break:
        doc.add_page_break()
    doc.add_heading(f"第{number}章 {title}", level=1)
    add_kv_table(doc, [("担当", owner), ("実施の目安", timing), ("この章の目的", purpose)])
    for heading, items in blocks:
        doc.add_heading(heading, level=2)
        if isinstance(items, str):
            add_text(doc, items)
        else:
            add_bullets(doc, items)
    if placeholder:
        add_placeholder(doc, placeholder)
    add_note(doc, "チェックポイント", " / ".join(checkpoints), "confirm")
    if faq:
        doc.add_heading("よくある質問", level=2)
        add_matrix(doc, ["質問", "一次対応"], faq, [2.2, 4.4])


def add_checklist_page(doc, title, items, owner="確認担当: ____________________", date_label="確認日: ____________________"):
    doc.add_page_break()
    doc.add_heading(title, level=1)
    add_text(doc, f"{owner}    {date_label}", color="475467", size=9.5)
    rows = [["□", item, "確認結果・補足"] for item in items]
    add_matrix(doc, ["", "確認項目", "記録"], rows, [0.35, 3.75, 2.5])
    add_note(doc, "記録", "未完了項目は担当者・期限・対応内容を記載し、次年度へ引き継いでください。", "notice")


def build_manual(filename, manual_name, subtitle, entries, chapters, checklists, force_break=True, compact=False):
    doc = Document()
    setup_document(doc, manual_name)
    if compact:
        # Keep the requested reference-guide page range without removing content.
        section = doc.sections[0]
        section.top_margin = Cm(1.55)
        section.bottom_margin = Cm(1.45)
        normal = doc.styles["Normal"]
        normal.font.size = Pt(9.7)
        normal.paragraph_format.space_after = Pt(3.5)
        normal.paragraph_format.line_spacing = 1.1
    add_cover(doc, manual_name, subtitle)
    add_revision_and_toc(doc, entries)
    for chapter in chapters:
        add_chapter(doc, *chapter, force_break=force_break)
    for checklist in checklists:
        add_checklist_page(doc, *checklist)
    doc.save(OUT / filename)


ADMIN_CHAPTERS = [
    (1, "マニュアルの目的", "理事長・専務・事務局", "就任時、年度初め", "会員情報と年度運営情報を安全に継続管理する。", [
        ("システムの目的", ["会員、年度、委員会、予定、出欠、お知らせ、資料を一元管理する。", "会員基本情報と年度ごとの役職・委員会・権限を分けて扱う。"]),
        ("管理者の責任", ["登録内容を本人確認のうえ更新する。", "個人情報を必要以上に出力・共有しない。", "パスワードを他人と共有しない。"]),
    ], "ダッシュボードと管理メニュー", ["管理責任者を年度ごとに決めた", "連絡先と保管場所を確認した"], None),
    (2, "管理者権限", "理事長・専務・年度管理担当", "年度切替前・就任直後", "誰が管理操作できるかを年度ごとに確認する。", [
        ("権限の条件", ["現在年度かつ有効な admin / president / secretary のみ管理操作を実行できる。", "役職・権限は年度所属管理で設定する。"]),
        ("運用原則", ["管理者は必要最小限にする。", "役職変更後は旧担当者の有効状態を確認する。", "一般会員で管理APIが拒否されることを年度ごとに確認する。"]),
    ], "年度所属編集画面", ["新年度の管理者を有効化した", "旧年度のみの権限を見直した"], None),
    (3, "ログイン・ログアウト", "全管理者", "日常", "安全にログインし、端末利用後にログアウトする。", [
        ("ログイン", ["ログインIDは会員メールアドレス。", "初回は管理者から案内された初期パスワードでログインし、直ちに変更する。", "変更後は8文字以上で英字と数字をそれぞれ1文字以上含める。"]),
        ("トラブル時", ["失敗表示はメール存在有無を示さない。", "5回失敗時の30秒停止は画面上の補助であり、管理者は認証防御として依存しない。", "初期パスワード再発行は本人確認後に行う。"]),
    ], "ログイン画面と初回変更画面", ["共用端末でログアウトした", "再発行前に本人確認をした"], [("ログインできない", "メールアドレス、入力文字種、管理者からの案内を確認する。")]),
    (4, "会員管理", "事務局・年度管理担当", "入会、異動、退会時", "会員基本情報とAuth連携を正確に管理する。", [
        ("新規登録", ["氏名、フリガナ、メールアドレス、電話番号、入会年度、ステータスを確認する。", "メールは重複登録しない。", "アカウント同時発行を標準とし、初期情報は本人へ安全な方法で渡す。"]),
        ("状態変更", ["休会・卒業・退会は削除ではなくステータス変更を基本とする。", "復帰時は会員情報と年度所属を再確認する。", "Auth未連携会員は初期パスワード発行または補助の招待メールを使う。"]),
    ], "会員一覧、会員登録、会員詳細", ["メール重複を確認した", "初期情報を画面表示後に安全に渡した", "退会者を物理削除していない"], [("初期情報を失くした", "本人確認後、会員詳細から初期パスワードを再発行する。")]),
    (5, "年度管理", "理事長・専務・年度管理担当", "年度末から年度初め", "毎年変わる役職・委員会を過年度データと混同せず管理する。", [
        ("新年度作成", ["年度名、開始日、終了日を確認して作成する。", "前年度コピー後は、役職・委員会・管理権限を必ず見直す。"]),
        ("年度切替", ["現在年度を変更する前に管理者と予定の年度を確認する。", "過年度データは削除せず参照用に残す。", "誤切替時は影響範囲を記録し、システム担当へ連絡する。"]),
    ], "年度一覧と年度詳細", ["開始日・終了日を確認した", "コピー後の権限を確認した"], None),
    (6, "委員会・役職管理", "専務・事務局・年度管理担当", "年度作成後、異動時", "年度ごとの組織体制を反映する。", [
        ("委員会", ["委員会名、説明、担当副理事長、委員長、副委員長を設定する。", "1会員は複数委員会を兼任できる。", "主所属は1つを目安に設定する。"]),
        ("役職と権限", ["年度所属で役職、権限、有効・無効を設定する。", "管理権限に直結する役職は、就任者・退任者を二重確認する。"]),
    ], "委員会詳細と年度所属編集", ["兼任を登録した", "主所属を確認した", "管理者権限を二重確認した"], None),
    (7, "予定管理", "事業担当・委員会担当", "予定確定時、変更時", "対象者と運営情報を含めた予定を登録する。", [
        ("作成", ["年度、種別、日時、会場、住所、説明を入力する。", "対象委員会、役職、個別会員を設定する。", "出欠が必要な場合は期限を設定する。"]),
        ("変更・削除", ["日時・会場変更時は対象者と出欠期限を見直す。", "削除は論理削除を前提とし、必要に応じてお知らせで周知する。", "年度をまたぐ予定は対象年度を誤らない。"]),
    ], "予定一覧、予定作成、予定詳細", ["対象者を確認した", "出欠期限を確認した", "変更を周知した"], None),
    (8, "出欠管理", "事業担当・事務局", "期限前、事業後", "回答状況を確認し、必要な催促と記録を行う。", [
        ("日常確認", ["回答率、未回答者、期限超過を確認する。", "委員会別・役職別の集計を確認する。", "締切後の変更は理由を記録して担当者間で共有する。"]),
        ("訂正", ["本人以外の回答変更は、本人確認または担当者の根拠を残す。", "CSV出力の利用可否・出力内容は運用開始前に要確認。"]),
    ], "出欠一覧、回答画面、集計画面", ["未回答者を確認した", "締切超過を確認した", "訂正理由を残した"], [("回答できない", "対象者設定、期限、ログイン会員との紐付けを確認する。")]),
    (9, "お知らせ管理", "理事長・専務・事務局", "連絡事項発生時", "対象者に必要な連絡を、重要度と公開期間を含めて発信する。", [
        ("投稿", ["種別、対象年度、対象委員会、公開範囲、重要度、公開期間を設定する。", "至急・重要は内容と対象を二重確認する。"]),
        ("編集・削除", ["誤記訂正は更新日時が残ることを前提に行う。", "公開終了後の扱いと削除可否は年度内のルールに従う。", "外部プッシュ通知の配信は要確認。"]),
    ], "お知らせ一覧、作成、詳細", ["公開対象を確認した", "重要度を確認した", "個人情報を含めていない"], None),
    (10, "資料管理", "事務局・各委員会担当", "資料配布時", "年度・イベント・公開範囲に沿って資料を整理する。", [
        ("登録", ["タイトル、ファイル名、カテゴリ、年度、関連イベント、公開範囲を確認する。", "ファイル名は「年度_委員会_資料名_版数」などの規則で統一する。"]),
        ("取扱い", ["個人情報・未公開資料は公開範囲を限定する。", "差し替え時は旧版の扱いを記録する。", "実ファイルのSupabase Storage保存は運用開始前に要確認。"]),
    ], "資料一覧、登録、詳細", ["公開範囲を確認した", "ファイル名ルールに従った", "実ファイルの保存先を確認した"], None),
    (11, "日常運用", "事務局・年度管理担当", "毎日・毎週・毎月", "未対応事項をためず、年度情報を最新に保つ。", [
        ("毎日", ["今日・今週の予定、未回答出欠、重要お知らせを確認する。"]),
        ("毎週・毎月", ["期限超過出欠、会員ステータス、予定変更、資料公開範囲を確認する。", "月次で管理者一覧と監査ログの要確認事項を見直す。"]),
        ("イベント発生時", ["新入会員は会員登録、年度所属、初期情報案内まで実施する。", "退会者はステータス、年度所属、公開対象を見直す。"]),
    ], "ダッシュボード", ["未回答を確認した", "今週の締切を確認した", "会員異動を反映した"], None),
    (12, "トラブル対応", "事務局一次対応・システム担当", "障害・問い合わせ時", "個人情報や認証情報を漏らさず、一次切り分けを行う。", [
        ("一次確認", ["URL、ログイン状態、対象年度、対象LOM、ブラウザ再読み込みを確認する。", "Application error、RLS拒否、Supabaseエラーは画面・時刻・操作手順だけを記録する。"]),
        ("エスカレーション", ["Secret key、初期パスワード、認証トークンを送らない。", "管理者で解決できない場合はシステム担当へ、再現手順と画面名を連絡する。"]),
    ], "エラー表示の例（秘密情報を写さない）", ["秘密情報を記録していない", "再現手順を記録した", "システム担当へ連絡した"], [("権限がない", "現在年度の役職、有効状態、ログイン会員の紐付けを確認する。"), ("招待メールが届かない", "メールアドレス、Supabase設定、送信制限を確認する。")]),
    (13, "管理者チェックリスト", "年度管理担当", "月次・年度切替前", "重要な操作を担当者間で相互確認する。", [
        ("使い方", ["次頁以降のチェックシートを印刷またはPDF注釈で使用する。", "未完了項目は担当者・期限・対応内容を残す。"]),
        ("必須記録", ["新規会員の発行日、再発行日、年度切替日、権限確認日を記録する。"]),
    ], "チェックシート記入例", ["担当者を記入した", "未完了項目の期限を設定した"], None),
]


YEAR_CHAPTERS = [
    (1, "年度引継ぎの目的", "現年度・次年度理事長、専務、事務局", "年度末の3か月前", "組織交代後もデータを消さず、責任と権限を確実に引き継ぐ。", [
        ("基本原則", ["会員基本情報は継続利用し、役職・委員会・権限は年度ごとに更新する。", "引継ぎ責任者と最終承認者を年度内に決める。"]),
    ], "年度一覧", ["現年度・次年度の責任者を決めた"], None),
    (2, "年度末の準備", "現年度専務・事務局", "年度末の1〜2か月前", "未処理データと管理者情報を整理する。", [
        ("確認項目", ["会員一覧、委員会、役職、未回答出欠、資料、お知らせを確認する。", "管理者一覧、GitHub、Supabaseの保管・連絡先を確認する。", "バックアップ手順と復元担当を確認する。"]),
    ], "年度末確認用ダッシュボード", ["未回答を整理した", "資料・お知らせを整理した", "バックアップ担当を確認した"], None),
    (3, "新年度作成", "次年度専務・年度管理担当", "年度開始前", "前年度を参照しながら、新年度の器を作成する。", [
        ("手順", ["年度名、開始日、終了日を入力して新年度を作成する。", "必要に応じて前年度コピーを使う。", "コピー後に役職、委員会、予定、管理権限を必ず見直す。"]),
        ("注意", ["コピー内容がそのまま正しいとは限らない。", "現在年度の切替は、次章以降の設定確認後に行う。"]),
    ], "新年度作成・前年度コピー", ["日付を確認した", "コピー後の内容を見直した"], None),
    (4, "会員整理", "事務局・年度管理担当", "新年度作成後", "継続会員と異動会員を正しく区分する。", [
        ("区分", ["卒業・退会・休会・継続・新入会員を確認する。", "メールアドレス変更時は会員情報とログインIDの扱いを確認する。", "既存Auth連携をむやみに上書きしない。"]),
    ], "会員一覧と会員詳細", ["卒業・退会を反映した", "新入会員を登録した", "メール変更を確認した"], None),
    (5, "新役職・委員会設定", "次年度理事長・専務", "新年度開始前", "新体制と兼任を年度所属へ反映する。", [
        ("設定順", ["委員会と役職を作成する。", "会員ごとに年度所属、役職、権限、有効状態を設定する。", "複数委員会、主所属、委員長・副委員長を確認する。"]),
    ], "委員会・年度所属編集", ["理事長・専務を確認した", "委員長・副委員長を確認した", "兼任・主所属を確認した"], None),
    (6, "管理者権限の引継ぎ", "現年度・次年度理事長、専務", "切替直前", "新年度の管理者だけが適切に管理できる状態にする。", [
        ("確認", ["次年度の admin / president / secretary を有効にする。", "現年度のみの管理者が新年度で不要なら有効ではないことを確認する。", "管理者と一般会員の両方でログイン・権限を確認する。"]),
    ], "年度所属と権限表示", ["新年度管理者を確認した", "一般会員で管理操作不可を確認した"], None),
    (7, "動作確認", "次年度事務局・システム担当", "切替前後", "主要画面とスマートフォン表示を確認する。", [
        ("確認画面", ["ホーム、会員、年度、委員会、予定、出欠、お知らせ、資料を確認する。", "スマートフォン幅でボタン・入力欄・文字が見切れないことを確認する。", "初期パスワード発行と初回変更をテスト会員で確認する。"]),
    ], "スマートフォン表示の確認位置", ["主要画面を確認した", "初期発行を確認した", "権限テストをした"], None),
    (8, "引継ぎ完了", "現年度・次年度担当者", "年度開始後1週間以内", "完了承認と未解決事項の記録を残す。", [
        ("完了条件", ["チェックシートの未完了項目を確認する。", "旧年度担当者、次年度担当者、確認日を記録する。", "不具合と要確認事項はシステム担当へ引き継ぐ。"]),
    ], "引継ぎ完了承認欄", ["承認者を記録した", "未解決事項を記録した"], None),
]


MEMBER_CHAPTERS = [
    (1, "利用開始", "会員本人", "初回ログイン時", "初期情報を受け取り、自分だけが使えるパスワードへ変更する。", [
        ("はじめに", ["管理者からログインIDと初期パスワードの案内を受ける。", "ログインIDは登録済みのメールアドレス。", "初回ログイン後はパスワード変更画面が開く。"]),
        ("パスワード", ["8文字以上で、英字と数字をそれぞれ1文字以上含める。", "パスワードを他人へ教えない。"]),
    ], "ログイン画面と初回変更画面", ["ログインIDを確認した", "初回変更を完了した"], None),
    (2, "ホーム画面", "会員本人", "毎日", "今日やることと未対応事項を確認する。", [
        ("見る場所", ["今日・今週の予定", "未回答の出欠", "お知らせ、新着資料", "自分の所属・役職"]),
    ], "ホーム画面", ["未回答出欠を確認した"], None),
    (3, "予定確認", "会員本人", "予定が近づいたとき", "日時、場所、対象者、持参物を確認する。", [
        ("操作", ["1. 下部メニュー等から予定を開く。", "2. 一覧から予定を選ぶ。", "3. 日時、場所、説明、出欠期限を確認する。"]),
    ], "予定一覧と予定詳細", ["日時・場所を確認した"], None),
    (4, "出欠回答", "会員本人", "期限前", "出席・欠席・遅刻を期限までに回答する。", [
        ("操作", ["1. 出欠対象の予定を開く。", "2. 出席、欠席、遅刻のいずれかを選ぶ。", "3. 必要ならコメントを入れて送信する。"]),
        ("注意", ["期限後の変更は担当者へ連絡する。", "回答内容を変更したい場合は画面の表示を確認し、できないときは事務局へ連絡する。"]),
    ], "出欠回答画面", ["回答を送信した", "期限を確認した"], None),
    (5, "お知らせ・資料", "会員本人", "連絡受信時", "対象のお知らせと資料を確認する。", [
        ("操作", ["1. お知らせから重要度と公開期間を確認する。", "2. 資料から対象のファイルを開く。", "3. 開けない資料は事務局へ資料名と画面名を連絡する。"]),
    ], "お知らせ・資料一覧", ["重要なお知らせを確認した"], None),
    (6, "自分の会員情報", "会員本人", "年度初め・変更時", "表示される氏名、所属、役職を確認する。", [
        ("確認", ["氏名、メールアドレス、所属、役職に誤りがある場合は事務局へ連絡する。", "自分で変更できる範囲は画面表示に従う。"]),
    ], "会員詳細", ["所属・役職を確認した"], None),
    (7, "ログアウト・パスワード", "会員本人", "共用端末利用後・忘失時", "共用端末にログイン状態を残さない。", [
        ("操作", ["1. 利用後はログアウトする。", "2. パスワードを忘れた場合は事務局へ連絡する。", "3. 本人確認後、管理者が初期パスワードを再発行する。", "4. 再発行後は再び初回変更を行う。"]),
    ], "ログアウト位置", ["共用端末でログアウトした"], None),
    (8, "よくある質問", "会員本人", "困ったとき", "画面上で解決しないときは、必要な情報だけを事務局へ伝える。", [
        ("連絡時に伝えること", ["画面名、操作した時刻、表示された文言。", "パスワードや認証コードは送らない。"]),
    ], None, ["秘密情報を送らない"], [("ログインできない", "メールアドレスと入力内容を確認し、事務局へ連絡する。"), ("予定が見えない", "対象者設定や年度所属の確認を事務局へ依頼する。"), ("資料が開かない", "資料名と画面名を事務局へ伝える。")]),
]


SYSTEM_CHAPTERS = [
    (1, "システム構成", "システム管理者", "引継ぎ時・構成変更時", "アプリ、認証、データ、ソース管理の責任分界を把握する。", [
        ("構成要素", ["Next.js / TypeScript: Webアプリ。", "Supabase Auth: メール・パスワード認証。", "Supabase Database: 会員・年度・出欠等のデータ。", "GitHub: ソースコードと変更履歴。", "Vercel: 本番ホスティング想定。実際のプロジェクト設定は要確認。"]),
        ("権限制御", ["RLSでLOM・本人・管理権限を制御する。", "管理権限は現在年度かつ有効な admin / president / secretary。"]),
    ], "システム構成図（差し替え用）", ["責任者と管理画面を確認した"], None),
    (2, "ローカル開発", "開発・保守担当", "変更前", "安全なローカル環境で変更を確認する。", [
        ("基本手順", ["リポジトリを取得する。", "依存関係をインストールする。", ".env.localへ公開URL、公開キー、サイトURL、サーバー専用Secret keyを設定する。", "開発サーバーを起動し、ポートを確認する。"]),
        ("禁止", [".env.local、Secret key、ログ、build出力、node_modulesをGitHubへ含めない。"]),
    ], "ローカル起動画面", ["環境変数を安全な保管場所で確認した", "秘密情報をコミット対象に含めていない"], None),
    (3, "GitHub運用", "開発・保守担当", "変更ごと", "変更を追跡可能な単位で記録する。", [
        ("基本", ["mainは安定版。feature/、fix/、docs/のブランチを使う。", "変更後は型チェック、lint、本番ビルドを実行する。", "コミットは短い命令形で記録する。"]),
        ("緊急時", ["Secretを誤ってPushした場合は直ちにキーを失効・再発行し、履歴だけで済ませない。", "リリースタグと復旧手順は要確認。"]),
    ], "GitHubリポジトリ画面", ["差分を確認した", "秘密情報がない", "検査結果を記録した"], None),
    (4, "Supabase運用", "システム管理者", "設定変更時", "認証、Database、RLS、URL設定を安全に扱う。", [
        ("確認場所", ["Authentication: ユーザー、URL Configuration、Password Security、Rate Limits、Bot and Abuse Protection。", "Database: Tables、SQL Editor、RLS policies。", "Project Settings: API Keys。Secret keyはサーバー専用。"]),
        ("注意", ["Publishable keyとSecret keyを混同しない。", "初期パスワード、認証トークン、会員の実メールを作業記録に残さない。"]),
    ], "Supabase Dashboardの設定位置", ["URL設定を確認した", "Password Securityを確認した", "Rate LimitsとCAPTCHAを確認した"], None),
    (5, "Migration", "システム管理者", "DB変更前", "SQLの目的と影響を確認してから反映する。", [
        ("実行前", ["対象ファイル、適用順、既存環境との差分を確認する。", "DROP、TRUNCATE、DELETEを含むSQLは必ず二重確認する。", "バックアップとロールバック方針を準備する。"]),
        ("実行後", ["SQL Editorの結果を確認する。", "主要画面、RLS、Auth連携を確認する。", "production-rls.sqlは再実行可能な最新版を使う。"]),
    ], "SQL Editor実行位置", ["適用順を確認した", "破壊的SQLを確認した", "反映後テストをした"], None),
    (6, "デプロイ", "システム管理者", "本番反映時", "本番URLと認証リダイレクトを整合させる。", [
        ("確認", ["Vercelの環境変数に公開URL、公開キー、サイトURL、Secret keyを設定する。", "Supabase Redirect URLsへ本番URLの認証コールバックを追加する。", "デプロイ後にログイン、管理操作、一般会員、未ログインを確認する。"]),
        ("要確認", ["Vercelプロジェクトの実在・権限・本番URL・ロールバック手順は引継ぎ時に確認する。"]),
    ], "デプロイ設定画面（要確認）", ["本番URLを確認した", "Redirect URLsを確認した", "ロールバックを確認した"], None),
    (7, "バックアップ", "システム管理者", "月次・変更前", "ソースとデータを復元できる状態にする。", [
        ("対象", ["GitHubのmainとリリース履歴。", "Supabase Databaseのエクスポートまたはバックアップ。", "重要な設定変更の記録。"]),
        ("運用", ["月次バックアップの実施日・担当者を記録する。", "年1回以上、復元手順を別環境で確認する。"]),
    ], "バックアップ記録欄", ["最新バックアップを確認した", "復元手順を確認した"], None),
    (8, "セキュリティ", "システム管理者", "定期確認", "秘密情報、RLS、管理API、監査ログを保護する。", [
        ("認証", ["Secret keyはRoute Handlerなどサーバー専用で使用する。", "管理APIは requireManagement() を通し、現在年度の有効な管理者だけを許可する。", "初期パスワード状態は通常画面をブロックする。"]),
        ("不正アクセス対策", ["画面内の30秒停止は補助表示であり、再読み込み等で回避可能。", "Supabase Rate LimitsとCAPTCHAを主防御とする。", "独自サーバー制限を導入する場合はメールとIPの組合せ、期限、成功時リセットを設計する。"]),
    ], "SECURITY.md確認位置", ["Secretを露出していない", "RLSを確認した", "CAPTCHA導入状況を確認した"], None),
    (9, "障害対応", "システム担当・事務局一次対応", "障害時", "秘密情報を漏らさず、再現可能な情報を収集する。", [
        ("代表例", ["Application error: ブラウザ、時刻、URL、直前操作を確認する。", "Invalid API key: 環境変数名とサーバー再起動を確認する。", "RLS拒否: ログイン会員、年度役職、LOM、policyを確認する。", "Auth紐付け不良: auth_user_idと会員レコードを管理手順に沿って確認する。"]),
        ("注意", ["fallback表示は本番の正常データとして扱わない。", "メール送信制限、Supabase障害、デプロイ失敗は公式ステータスと設定を確認する。"]),
    ], "障害記録テンプレート", ["秘密情報を記録していない", "再現手順を記録した"], None),
    (10, "更新・保守", "開発・保守担当", "機能変更時", "小さく変更し、検査と本番確認を残す。", [
        ("標準手順", ["変更前にバックアップと影響範囲を確認する。", "型チェック、lint、本番ビルドを実行する。", "管理者、一般会員、未ログインで確認する。", "GitHubへPushし、本番反映後に確認する。"]),
        ("Codexへの依頼", ["目的、対象画面、既存機能を壊さない条件、確認項目を具体的に記載する。", "Secretや実パスワードを会話へ貼り付けない。"]),
    ], "更新前後確認表", ["検査を実行した", "権限別テストをした", "変更履歴を残した"], None),
    (11, "引継ぎ情報", "現任・次年度システム担当", "交代前", "アカウント、契約、連絡先、保管場所を安全に引き継ぐ。", [
        ("引継ぐもの", ["各サービスの管理者、契約情報、更新期限、緊急連絡先。", "安全なパスワード管理ツール内の保管場所。", "GitHub、Supabase、Vercelの権限一覧。"]),
        ("書かないもの", ["Secret key、アクセストークン、実際のパスワードを引継ぎ文書へ書かない。"]),
    ], "引継ぎ情報記入欄", ["新旧担当者を確認した", "権限移管を確認した", "緊急連絡先を確認した"], None),
]


def build_checklists():
    doc = Document()
    setup_document(doc, "運用チェックリスト")
    add_cover(doc, "運用チェックリスト", "日常運用、年度引継ぎ、登録、予定・出欠、トラブル一次切り分け")
    add_checklist_page(doc, "1. 管理者向け日常運用チェックリスト", [
        "今日・今週の予定を確認した", "未回答・期限超過の出欠を確認した", "重要・至急のお知らせを確認した", "新着資料の公開範囲を確認した", "会員異動・予定変更を反映した", "未解決の問い合わせを記録した",
    ])
    add_checklist_page(doc, "2. 年度引継ぎチェックリスト（準備）", [
        "現年度会員一覧とステータスを確認した", "卒業・退会・休会・継続を確認した", "委員会・役職・兼任・主所属を確認した", "未回答出欠、資料、お知らせを整理した", "管理者一覧とバックアップ担当を確認した", "新旧担当者の連絡先を確認した",
    ], "確認担当: ____________________", "年度: ____________________")
    add_checklist_page(doc, "3. 年度引継ぎチェックリスト（切替・確認）", [
        "新年度を作成し、日付を確認した", "前年度コピー後の内容を修正した", "新年度の管理者を有効化した", "旧年度のみの管理権限を確認した", "管理者・一般会員で権限を確認した", "スマートフォンで主要画面を確認した", "引継ぎ完了日と未解決事項を記録した",
    ], "確認担当: ____________________", "年度: ____________________")
    doc.add_page_break()
    doc.add_heading("4. 新入会員登録フロー図", level=1)
    add_text(doc, "初期情報は本人確認後に安全な方法で渡し、文書やチャットへ残さない。", color="475467", size=9.5)
    add_matrix(doc, ["手順", "実施者", "確認"], [
        ["1. 会員情報を確認", "事務局", "氏名、メール、入会年度、ステータス"],
        ["2. 会員登録・アカウント同時発行", "管理者", "メール重複なし、初回変更必須"],
        ["3. 初期情報を本人へ案内", "管理者", "安全な伝達、画面を閉じた後は再表示不可"],
        ["4. 初回ログイン・変更", "会員本人", "新パスワード条件を満たす"],
        ["5. 年度所属を設定", "年度管理担当", "委員会、役職、権限、有効状態"],
        ["6. 利用開始確認", "事務局", "ホーム、予定、出欠の表示"],
    ], [1.65, 1.35, 3.3])
    add_placeholder(doc, "新規会員登録から初回変更までの画面キャプチャ")
    doc.add_page_break()
    doc.add_heading("5. 予定登録から出欠集計までのフロー図", level=1)
    add_matrix(doc, ["手順", "実施者", "確認"], [
        ["1. 予定作成", "事業担当", "年度、日時、会場、対象者"],
        ["2. 出欠対象を生成", "システム / 担当者", "対象委員会、役職、個別会員"],
        ["3. お知らせ・リマインド", "担当者", "期限、対象、変更周知"],
        ["4. 会員が回答", "会員本人", "出席、欠席、遅刻、コメント"],
        ["5. 締切確認", "事務局", "未回答、期限超過"],
        ["6. 集計・記録", "事業担当", "出席数、委員会別、役職別、訂正理由"],
    ], [1.65, 1.35, 3.3])
    add_placeholder(doc, "予定詳細、出欠回答、集計画面のキャプチャ")
    doc.add_page_break()
    doc.add_heading("6. トラブル一次切り分け表", level=1)
    add_matrix(doc, ["症状", "まず確認すること", "次の連絡先"], [
        ["ログインできない", "メール、入力内容、初回変更の有無", "事務局"],
        ["権限がない", "現在年度の役職・有効状態・LOM", "年度管理担当"],
        ["会員・予定が見えない", "対象年度、対象者設定、ログイン状態", "事務局"],
        ["出欠回答できない", "対象者、期限、イベント設定", "事業担当"],
        ["Application error", "URL、時刻、直前操作。秘密情報は送らない", "システム担当"],
        ["Supabase・RLSエラー", "設定変更の有無、エラー文、再現手順", "システム担当"],
    ], [1.55, 3.0, 1.75])
    add_note(doc, "禁止", "Secret key、初期パスワード、認証トークン、実会員メールを問い合わせ記録へ書かない。", "warning")
    doc.save(OUT / "05_運用チェックリスト.docx")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_manual(
        "01_管理者運用マニュアル.docx",
        "管理者運用マニュアル",
        "理事長・専務・事務局・年度管理担当者向け",
        [f"第{i}章 {chapter[1]}" for i, chapter in enumerate(ADMIN_CHAPTERS, 1)],
        ADMIN_CHAPTERS,
        [
            ("付録A 新規会員登録チェック", ["本人情報を確認した", "メール重複を確認した", "初期情報を安全に渡した", "年度所属を設定した"]),
            ("付録B 初期パスワード再発行チェック", ["本人確認をした", "再発行理由を確認した", "初回変更必須を案内した", "監査記録を確認した"]),
            ("付録C 月次確認チェック", ["未回答出欠を確認した", "管理者一覧を確認した", "予定・お知らせ・資料を確認した", "要確認事項を記録した"]),
            ("付録D 年度切替前チェック", ["次年度を作成した", "役職・委員会を確認した", "新年度管理者を確認した", "動作確認を実施した"]),
        ],
    )
    build_manual(
        "02_年度引継ぎマニュアル.docx",
        "年度引継ぎマニュアル",
        "現年度・次年度の理事長、専務、事務局、システム担当向け",
        [f"第{i}章 {chapter[1]}" for i, chapter in enumerate(YEAR_CHAPTERS, 1)],
        YEAR_CHAPTERS,
        [
            ("付録A 年度引継ぎ完了承認シート", ["現年度担当者の確認", "次年度担当者の確認", "システム担当の確認", "未解決事項と期限の記録"]),
            ("付録B 引継ぎ時の権限テスト", ["管理者で会員・年度を管理できる", "一般会員で管理操作が拒否される", "未ログインで保護画面がログインへ移動する", "スマートフォン表示を確認した"]),
            ("付録C 年度切替後1週間チェック", ["予定と出欠を確認した", "お知らせ・資料を確認した", "新入会員の初期発行を確認した", "問い合わせを記録した"]),
        ],
        compact=True,
    )
    build_manual(
        "03_一般会員利用マニュアル.docx",
        "一般会員利用マニュアル",
        "スマートフォン操作に慣れていない会員向け",
        [f"第{i}章 {chapter[1]}" for i, chapter in enumerate(MEMBER_CHAPTERS, 1)],
        MEMBER_CHAPTERS,
        [],
        force_break=False,
        compact=True,
    )
    build_manual(
        "04_システム管理者マニュアル.docx",
        "システム管理者マニュアル",
        "保守・開発・次年度引継ぎ担当者向け",
        [f"第{i}章 {chapter[1]}" for i, chapter in enumerate(SYSTEM_CHAPTERS, 1)],
        SYSTEM_CHAPTERS,
        [
            ("付録A リリース前チェック", ["型チェック", "lint", "本番ビルド", "管理者・一般会員・未ログインの確認", "環境変数の確認"]),
            ("付録B Migration実行記録", ["対象SQL", "実行者", "実行日時", "実行結果", "影響確認", "ロールバック可否"]),
            ("付録C セキュリティ定期確認", ["Secret keyの保管場所", "RLS", "管理者一覧", "Rate Limits", "CAPTCHA", "バックアップ"]),
        ],
    )
    build_checklists()


if __name__ == "__main__":
    main()
